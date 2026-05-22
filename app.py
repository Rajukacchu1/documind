import streamlit as st
import os
import re
import json
from pathlib import Path
import base64
from io import BytesIO

# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="DocuMind – Document Intelligence",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Custom CSS ─────────────────────────────────────────────────────────────────
st.markdown("""
<style>
/* ── Fonts ── */
@import url('https://fonts.googleapis.com/css2?family=Cinzel:wght@400;600;700&family=Inter:wght@300;400;500;600&family=JetBrains+Mono:wght@400;500&display=swap');

/* ── Root vars ── */
:root {
    --bg:          #0a0a0f;
    --bg2:         #0d0d15;
    --surface:     #12121e;
    --surface2:    #1a1a2e;
    --surface3:    #1e1e35;
    --border:      #2a2a4a;
    --border2:     #3a3a6a;
    --gold:        #f0c060;
    --gold2:       #e8a030;
    --violet:      #9d6fff;
    --violet2:     #7b4fe0;
    --teal:        #40e0c0;
    --text:        #f0eeff;
    --text-muted:  #9090b8;
    --text-dim:    #6060a0;
    --success:     #40c880;
    --radius:      12px;
    --radius-lg:   18px;
}

/* ── Global ── */
html, body { background-color: var(--bg) !important; }
[class*="css"], .stApp {
    font-family: 'Inter', sans-serif !important;
    background-color: var(--bg) !important;
    color: var(--text) !important;
}

/* ── Decorative background pattern ── */
[data-testid="stAppViewContainer"] {
    background:
        radial-gradient(ellipse 80% 40% at 20% 0%, rgba(157,111,255,.10) 0%, transparent 60%),
        radial-gradient(ellipse 60% 30% at 80% 100%, rgba(64,224,192,.07) 0%, transparent 60%),
        var(--bg) !important;
}

/* ══════════════════════════════════
   SIDEBAR
══════════════════════════════════ */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #0e0e1c 0%, #10101f 100%) !important;
    border-right: 1px solid var(--border2) !important;
    box-shadow: 4px 0 24px rgba(0,0,0,.5) !important;
}

/* Force all sidebar text white */
[data-testid="stSidebar"],
[data-testid="stSidebar"] *:not(button):not(.sidebar-brand h1) {
    color: var(--text) !important;
}

/* Sidebar section labels */
[data-testid="stSidebar"] h4,
[data-testid="stSidebar"] h3,
[data-testid="stSidebar"] .stMarkdown p {
    color: var(--text) !important;
    font-weight: 600 !important;
    letter-spacing: .03em !important;
}

/* ── Sidebar brand ── */
.sidebar-brand {
    background: linear-gradient(135deg, #161630 0%, #1a1040 100%);
    border-bottom: 1px solid var(--border2);
    padding: 30px 20px 24px;
    margin: -1rem -1rem 1.5rem;
    position: relative;
    overflow: hidden;
}
.sidebar-brand::before {
    content: '';
    position: absolute; inset: 0;
    background: radial-gradient(ellipse 70% 80% at 50% -10%, rgba(157,111,255,.25), transparent);
    pointer-events: none;
}
.sidebar-brand h1 {
    font-family: 'Cinzel', serif !important;
    font-size: 1.55rem !important;
    font-weight: 700 !important;
    background: linear-gradient(90deg, var(--gold), var(--violet), var(--teal));
    -webkit-background-clip: text !important;
    -webkit-text-fill-color: transparent !important;
    margin: 0 0 6px !important;
    letter-spacing: .06em !important;
}
.sidebar-brand p {
    color: var(--text-muted) !important;
    font-size: 0.76rem !important;
    margin: 0 !important;
    letter-spacing: .08em !important;
    text-transform: uppercase !important;
}

/* ── Sidebar input fields (folder path, api key) ── */
[data-testid="stSidebar"] .stTextInput input {
    background: var(--surface2) !important;
    border: 1px solid var(--border2) !important;
    border-radius: var(--radius) !important;
    color: var(--text) !important;
    caret-color: var(--violet) !important;
    font-size: 0.88rem !important;
    padding: 10px 14px !important;
}
[data-testid="stSidebar"] .stTextInput input::placeholder { color: var(--text-dim) !important; opacity: 1 !important; }
[data-testid="stSidebar"] .stTextInput input:focus {
    border-color: var(--violet) !important;
    box-shadow: 0 0 0 3px rgba(157,111,255,.18) !important;
    outline: none !important;
}

/* ── File uploader ── */
[data-testid="stFileUploader"] {
    background: transparent !important;
}
[data-testid="stFileUploaderDropzone"] {
    background: var(--surface2) !important;
    border: 2px dashed var(--border2) !important;
    border-radius: var(--radius-lg) !important;
    transition: border-color .2s, background .2s !important;
}
[data-testid="stFileUploaderDropzone"]:hover {
    border-color: var(--violet) !important;
    background: var(--surface3) !important;
}
/* all text inside the dropzone */
[data-testid="stFileUploaderDropzone"] *,
[data-testid="stFileUploaderDropzone"] span,
[data-testid="stFileUploaderDropzone"] small,
[data-testid="stFileUploaderDropzone"] p,
[data-testid="stFileUploaderDropzone"] div,
[data-testid="stFileUploaderDropzoneInstructions"] *,
[data-testid="stFileUploaderDropzoneInstructions"] span,
[data-testid="stFileUploaderDropzoneInstructions"] small {
    color: var(--text) !important;
    opacity: 1 !important;
}
/* secondary hint text */
[data-testid="stFileUploaderDropzone"] small,
[data-testid="stFileUploaderDropzoneInstructions"] small {
    color: var(--text-muted) !important;
}
/* Browse files button */
[data-testid="stFileUploaderDropzone"] button,
[data-testid="stFileUploaderDropzone"] [data-testid="baseButton-secondary"] {
    background: linear-gradient(135deg, var(--violet2), var(--violet)) !important;
    color: #fff !important;
    border: none !important;
    border-radius: var(--radius) !important;
    font-weight: 600 !important;
    padding: 8px 20px !important;
    box-shadow: 0 2px 12px rgba(157,111,255,.35) !important;
}
/* uploaded file name pills */
[data-testid="stFileUploader"] [data-testid="stFileUploaderFile"] {
    background: rgba(157,111,255,.12) !important;
    border: 1px solid rgba(157,111,255,.3) !important;
    border-radius: var(--radius) !important;
    color: var(--text) !important;
}
[data-testid="stFileUploader"] [data-testid="stFileUploaderFile"] * {
    color: var(--text) !important;
}

/* ── Sidebar buttons ── */
[data-testid="stSidebar"] .stButton > button {
    background: linear-gradient(135deg, var(--violet2), var(--violet)) !important;
    color: #fff !important;
    border: none !important;
    border-radius: var(--radius) !important;
    font-family: 'Inter', sans-serif !important;
    font-weight: 600 !important;
    font-size: 0.85rem !important;
    letter-spacing: .04em !important;
    padding: 9px 16px !important;
    transition: all .2s !important;
    box-shadow: 0 2px 12px rgba(157,111,255,.3) !important;
}
[data-testid="stSidebar"] .stButton > button:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 6px 20px rgba(157,111,255,.45) !important;
}

/* ── Stats cards ── */
.stat-card {
    background: linear-gradient(135deg, var(--surface2), var(--surface3));
    border: 1px solid var(--border2);
    border-radius: var(--radius);
    padding: 14px 12px;
    text-align: center;
    box-shadow: 0 4px 16px rgba(0,0,0,.4);
}
.stat-card .num {
    font-family: 'Cinzel', serif;
    font-size: 1.8rem; font-weight: 700;
    background: linear-gradient(90deg, var(--gold), var(--violet));
    -webkit-background-clip: text; -webkit-text-fill-color: transparent;
}
.stat-card .lbl {
    color: var(--text-muted) !important;
    font-size: 0.7rem; letter-spacing: .08em; text-transform: uppercase; margin-top: 4px;
}

/* ══════════════════════════════════
   MAIN AREA
══════════════════════════════════ */
[data-testid="stMain"],
[data-testid="stMain"] * { color: var(--text) !important; }

/* ── Main header ── */
.main-header {
    padding: 2.5rem 1rem 1.8rem;
    text-align: center;
    margin-bottom: 1.5rem;
    position: relative;
}
.main-header::after {
    content: '';
    display: block;
    width: 120px; height: 2px;
    background: linear-gradient(90deg, transparent, var(--gold), var(--violet), transparent);
    margin: 1rem auto 0;
    border-radius: 2px;
}
.main-header h2 {
    font-family: 'Cinzel', serif !important;
    font-size: 2.2rem !important; font-weight: 700 !important;
    background: linear-gradient(90deg, var(--gold), var(--violet), var(--teal));
    -webkit-background-clip: text !important; -webkit-text-fill-color: transparent !important;
    margin: 0 0 8px !important; letter-spacing: .05em !important;
}
.main-header p { color: var(--text-muted) !important; font-size: 0.88rem !important; margin: 0 !important; letter-spacing: .04em !important; }

/* ── Chat bubbles ── */
.chat-wrap { display: flex; flex-direction: column; gap: 1.4rem; padding-bottom: 1rem; }

.msg-user { display: flex; justify-content: flex-end; }
.msg-user .bubble {
    background: linear-gradient(135deg, var(--violet2) 0%, #5a2fb0 100%);
    color: #fff !important;
    padding: 14px 20px;
    border-radius: var(--radius-lg) var(--radius-lg) 4px var(--radius-lg);
    max-width: 70%; font-size: 0.93rem; line-height: 1.65;
    box-shadow: 0 6px 24px rgba(123,79,224,.35);
    border: 1px solid rgba(157,111,255,.3);
}

.msg-bot { display: flex; align-items: flex-start; gap: 12px; }
.msg-bot .avatar {
    width: 38px; height: 38px; border-radius: 50%; flex-shrink: 0; margin-top: 2px;
    background: linear-gradient(135deg, var(--gold2), var(--violet));
    display: flex; align-items: center; justify-content: center; font-size: 1.05rem;
    box-shadow: 0 4px 14px rgba(157,111,255,.4);
}
.msg-bot .bubble {
    background: linear-gradient(135deg, var(--surface) 0%, var(--surface2) 100%);
    border: 1px solid var(--border2);
    color: var(--text) !important;
    padding: 16px 22px;
    border-radius: 4px var(--radius-lg) var(--radius-lg) var(--radius-lg);
    max-width: 80%; font-size: 0.93rem; line-height: 1.75;
    box-shadow: 0 6px 28px rgba(0,0,0,.45);
}

/* ── Source badge ── */
.source-badge {
    display: inline-flex; align-items: center; gap: 6px;
    background: rgba(157,111,255,.12); border: 1px solid rgba(157,111,255,.3);
    border-radius: 20px; padding: 4px 12px;
    font-size: 0.72rem; color: var(--text-muted) !important;
    margin-top: 10px; margin-right: 6px;
    font-family: 'JetBrains Mono', monospace;
}

/* ── Query input (search bar) ── */
.stTextInput > div > div > input {
    background: var(--surface2) !important;
    border: 1.5px solid var(--border2) !important;
    border-radius: var(--radius-lg) !important;
    color: var(--text) !important;
    caret-color: var(--violet) !important;
    font-family: 'Inter', sans-serif !important;
    font-size: 0.95rem !important;
    padding: 14px 20px !important;
    transition: all .25s !important;
}
.stTextInput > div > div > input::placeholder { color: var(--text-dim) !important; opacity: 1 !important; }
.stTextInput > div > div > input:focus {
    border-color: var(--violet) !important;
    background: var(--surface3) !important;
    box-shadow: 0 0 0 4px rgba(157,111,255,.15), 0 4px 20px rgba(157,111,255,.12) !important;
    outline: none !important;
}

/* ── Main form submit button ── */
[data-testid="stForm"] .stButton > button,
.stFormSubmitButton > button {
    background: linear-gradient(135deg, var(--gold2) 0%, var(--gold) 100%) !important;
    color: #0a0a0f !important;
    border: none !important;
    border-radius: var(--radius-lg) !important;
    font-family: 'Inter', sans-serif !important;
    font-weight: 700 !important;
    font-size: 0.92rem !important;
    letter-spacing: .05em !important;
    padding: 14px 20px !important;
    transition: all .2s !important;
    box-shadow: 0 4px 18px rgba(240,192,96,.3) !important;
}
[data-testid="stForm"] .stButton > button:hover,
.stFormSubmitButton > button:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 8px 28px rgba(240,192,96,.45) !important;
}

/* ── Query form — inline, within normal page flow ── */
[data-testid="stForm"] {
    position: static !important;
    background: var(--surface2) !important;
    padding: 14px 1.5rem 12px !important;
    border: 1px solid var(--border2) !important;
    border-radius: var(--radius) !important;
    box-shadow: 0 2px 12px rgba(0,0,0,.35) !important;
    margin-top: 0.5rem !important;
}

/* ── Tables ── */
.msg-bot .bubble table { border-collapse: collapse; width: 100%; margin: 12px 0; font-size: 0.85rem; }
.msg-bot .bubble th {
    background: rgba(157,111,255,.15); color: var(--gold) !important;
    padding: 9px 14px; text-align: left; border: 1px solid var(--border2);
    font-family: 'JetBrains Mono', monospace; font-size: 0.78rem; letter-spacing: .05em;
}
.msg-bot .bubble td {
    padding: 8px 14px; border: 1px solid var(--border); vertical-align: top; color: var(--text) !important;
}
.msg-bot .bubble tr:nth-child(even) td { background: rgba(255,255,255,.03); }

/* ── Code blocks ── */
.msg-bot .bubble code, .msg-bot .bubble pre {
    background: rgba(0,0,0,.4) !important; border-radius: 8px;
    font-family: 'JetBrains Mono', monospace; font-size: 0.82rem;
    color: var(--teal) !important; border: 1px solid var(--border);
}

/* ── Divider ── */
hr { border-color: var(--border2) !important; }

/* ── Alerts ── */
.stAlert { border-radius: var(--radius) !important; }
.stAlert * { color: inherit !important; }

/* ── Scrollbar ── */
::-webkit-scrollbar { width: 7px; }
::-webkit-scrollbar-track { background: #ede9fe; border-radius: 4px; }
::-webkit-scrollbar-thumb { background: #7c3aed; border-radius: 4px; }
::-webkit-scrollbar-thumb:hover { background: #6d28d9; }

/* ── Dataframe toolbar ── */
/* wrapper — no overflow:hidden so toolbar overlay is never clipped */
[data-testid="stDataFrame"] { border-radius: var(--radius) !important; }

/* toolbar container */
[data-testid="stElementToolbar"] {
    background: var(--surface3) !important;
    border: 1px solid var(--border2) !important;
    border-radius: 10px !important;
    padding: 4px 6px !important;
    gap: 2px !important;
    opacity: 1 !important;
    visibility: visible !important;
    display: flex !important;
    align-items: center !important;
}

/* every button in the toolbar */
[data-testid="stElementToolbar"] button {
    background: var(--surface2) !important;
    border: 1px solid var(--border2) !important;
    border-radius: 7px !important;
    color: var(--text) !important;
    width: 30px !important;
    height: 30px !important;
    padding: 5px !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    cursor: pointer !important;
    opacity: 1 !important;
    visibility: visible !important;
    transition: background .15s, border-color .15s !important;
}
[data-testid="stElementToolbar"] button:hover {
    background: rgba(157,111,255,.25) !important;
    border-color: var(--violet) !important;
}

/* force every SVG shape inside toolbar buttons to be white */
[data-testid="stElementToolbar"] button svg { overflow: visible !important; }
[data-testid="stElementToolbar"] button svg *,
[data-testid="stElementToolbar"] button svg path,
[data-testid="stElementToolbar"] button svg polyline,
[data-testid="stElementToolbar"] button svg line,
[data-testid="stElementToolbar"] button svg rect,
[data-testid="stElementToolbar"] button svg circle,
[data-testid="stElementToolbar"] button svg polygon {
    stroke: var(--text) !important;
    fill: none !important;
    opacity: 1 !important;
    visibility: visible !important;
}
[data-testid="stElementToolbar"] button:hover svg *,
[data-testid="stElementToolbar"] button:hover svg path,
[data-testid="stElementToolbar"] button:hover svg polyline,
[data-testid="stElementToolbar"] button:hover svg line,
[data-testid="stElementToolbar"] button:hover svg rect,
[data-testid="stElementToolbar"] button:hover svg circle,
[data-testid="stElementToolbar"] button:hover svg polygon {
    stroke: var(--violet) !important;
}

/* ── Dataframe search input (appears in toolbar when search is clicked) ── */
[data-testid="stDataFrame"] input,
[data-testid="stDataFrameResizable"] input,
[data-testid="stElementToolbar"] input,
[data-testid="stElementToolbar"] + div input,
[class*="stDataFrame"] input,
[class*="dataframe"] input {
    background: var(--surface2) !important;
    color: var(--text) !important;
    border: 1px solid var(--border2) !important;
    border-radius: 8px !important;
    caret-color: var(--violet) !important;
    outline: none !important;
}
[data-testid="stDataFrame"] input::placeholder,
[data-testid="stDataFrameResizable"] input::placeholder,
[data-testid="stElementToolbar"] input::placeholder {
    color: var(--text-dim) !important;
    opacity: 1 !important;
}
[data-testid="stDataFrame"] input:focus,
[data-testid="stDataFrameResizable"] input:focus,
[data-testid="stElementToolbar"] input:focus {
    border-color: var(--violet) !important;
    box-shadow: 0 0 0 3px rgba(157,111,255,.2) !important;
}

/* ── Thinking / loading dots ── */
.thinking-dots {
    display: inline-flex;
    gap: 5px;
    align-items: center;
    height: 20px;
}
.thinking-dots span {
    width: 8px; height: 8px;
    border-radius: 50%;
    background: var(--violet);
    display: inline-block;
    animation: thinking-bounce 1.2s infinite ease-in-out;
}
.thinking-dots span:nth-child(1) { animation-delay: 0s; }
.thinking-dots span:nth-child(2) { animation-delay: 0.2s; }
.thinking-dots span:nth-child(3) { animation-delay: 0.4s; }
@keyframes thinking-bounce {
    0%, 80%, 100% { transform: scale(0.6); opacity: 0.4; }
    40%            { transform: scale(1.2); opacity: 1; }
}

/* ── Message action buttons (📋 Copy, 🔄 Reload) — compact rectangles ── */
button[data-testid="baseButton-secondary"] {
    padding: 2px 8px !important;
    min-height: 30px !important;
    height: 30px !important;
    font-size: .8rem !important;
    border-radius: 5px !important;
    background: #ede9fe !important;
    border: 1.5px solid #7c3aed !important;
    color: #3b0764 !important;
    font-weight: 600 !important;
    line-height: 1 !important;
    white-space: nowrap !important;
}
button[data-testid="baseButton-secondary"]:hover {
    background: #ddd6fe !important;
    border-color: #6d28d9 !important;
}

/* ── Stop button — styled via dynamic injection in Python, not here ── */

/* ── Hide Streamlit chrome ── */
#MainMenu, footer { visibility: hidden; }

/* ── Header — themed to match app ── */
header[data-testid="stHeader"] {
    background: linear-gradient(90deg, #0e0e1c 0%, #10101f 100%) !important;
    border-bottom: 1px solid var(--border2) !important;
    box-shadow: 0 2px 16px rgba(0,0,0,.4) !important;
}

/* All text/icons inside header */
header[data-testid="stHeader"] * {
    color: var(--text) !important;
    fill: var(--text) !important;
}

/* Deploy button */
header[data-testid="stHeader"] [data-testid="stDeployButton"],
header[data-testid="stHeader"] [data-testid="stDeployButton"] * {
    background: linear-gradient(135deg, var(--violet2), var(--violet)) !important;
    color: #fff !important;
    border: none !important;
    border-radius: var(--radius) !important;
    font-family: 'Inter', sans-serif !important;
    font-weight: 600 !important;
    font-size: .82rem !important;
    letter-spacing: .03em !important;
}
header[data-testid="stHeader"] [data-testid="stDeployButton"]:hover {
    background: linear-gradient(135deg, var(--violet), var(--gold)) !important;
    box-shadow: 0 0 12px rgba(157,111,255,.5) !important;
}

/* Toolbar icon buttons (fullscreen, share, etc.) */
header[data-testid="stHeader"] button {
    color: var(--text-muted) !important;
    background: transparent !important;
    border: none !important;
    border-radius: 8px !important;
    transition: background .2s, color .2s !important;
}
header[data-testid="stHeader"] button:hover {
    background: var(--surface3) !important;
    color: var(--gold) !important;
}
header[data-testid="stHeader"] button svg {
    fill: var(--text-muted) !important;
}
header[data-testid="stHeader"] button:hover svg {
    fill: var(--gold) !important;
}

.block-container { padding-top: 1rem !important; padding-bottom: 2rem !important; }
</style>
""", unsafe_allow_html=True)

# ── Imports (after page config) ───────────────────────────────────────────────
try:
    import fitz  # PyMuPDF
    PYMUPDF_OK = True
except ImportError:
    PYMUPDF_OK = False

try:
    from docx import Document as DocxDocument
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    import pandas as pd
    PANDAS_OK = True
except ImportError:
    PANDAS_OK = False

try:
    from sentence_transformers import SentenceTransformer
    import numpy as np
    EMBEDDING_OK = True
except ImportError:
    EMBEDDING_OK = False

try:
    import anthropic
    ANTHROPIC_OK = True
except ImportError:
    ANTHROPIC_OK = False

# ── Session state ─────────────────────────────────────────────────────────────
if "messages" not in st.session_state:
    st.session_state.messages = []
if "doc_chunks" not in st.session_state:
    st.session_state.doc_chunks = []   # list of {text, source, page, images, tables}
if "embeddings" not in st.session_state:
    st.session_state.embeddings = None
if "loaded_files" not in st.session_state:
    st.session_state.loaded_files = []
if "doc_pdf_bytes" not in st.session_state:
    st.session_state.doc_pdf_bytes = {}   # filename → raw PDF bytes for page rendering
if "page_img_cache" not in st.session_state:
    st.session_state.page_img_cache = {}  # (filename, page_num, mode) → base64 PNG
_LEARNED_FILE = Path(__file__).parent / "learned_answers.json"


def _save_learned_to_disk(learned: list) -> None:
    """Persist confirmed Q&A pairs to disk (images stripped to keep file small)."""
    saveable = []
    for la in learned:
        msg = {k: v for k, v in la["message"].items() if k != "images"}
        saveable.append({"query": la["query"], "message": msg})
    try:
        _LEARNED_FILE.write_text(
            json.dumps(saveable, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    except Exception:
        pass


def _load_learned_from_disk() -> list:
    """Load persisted Q&A pairs, restoring default fields."""
    if not _LEARNED_FILE.exists():
        return []
    try:
        data = json.loads(_LEARNED_FILE.read_text(encoding="utf-8"))
        for la in data:
            la["message"].setdefault("images", [])
            la["message"].setdefault("tables", [])
            la["message"].setdefault("sources", [])
        return data
    except Exception:
        return []


if "learned_answers" not in st.session_state:
    st.session_state.learned_answers = _load_learned_from_disk()

# ── Helpers ───────────────────────────────────────────────────────────────────

def img_to_b64(img_bytes: bytes) -> str:
    return base64.b64encode(img_bytes).decode()


def _get_media_type(img_b64: str) -> str:
    """Detect image MIME type from the opening bytes of its base64 string."""
    if img_b64.startswith("/9j/"):
        return "image/jpeg"
    if img_b64.startswith("R0lG"):
        return "image/gif"
    if img_b64.startswith("Qk"):
        return "image/bmp"
    return "image/png"


def analyze_image(img_b64: str, query: str, api_key: str):
    """
    Classify an image in the context of the user's query.
    Returns (kind, data) where kind is:
      "table"      → image is a table; data = list of rows
      "relevant"   → image is relevant to the query (diagram, figure, text); data = None
      "irrelevant" → logo, watermark, decoration; data = None
    Results are cached per (img_b64, query) so the same image isn't re-sent twice.
    """
    if not ANTHROPIC_OK or not api_key:
        return "irrelevant", None
    cache = st.session_state.setdefault("_img_cache", {})
    cache_key = (img_b64[:64], query[:120])
    if cache_key in cache:
        return cache[cache_key]
    try:
        client = anthropic.Anthropic(api_key=api_key)
        media_type = _get_media_type(img_b64)
        prompt = (
            f"The user asked: \"{query[:200]}\"\n\n"
            "Look at this image and respond with EXACTLY ONE of the following:\n\n"
            "TABLE\n[[row1col1, row1col2, ...], [row2col1, ...], ...]\n"
            "  → Use this if the image IS a table, data grid, or spreadsheet. "
            "    Extract every row and column as a JSON array of arrays.\n\n"
            "RELEVANT\n"
            "  → Use this if the image is a diagram, chart, figure, flowchart, or contains "
            "    text/data that is useful for answering the question.\n\n"
            "IRRELEVANT\n"
            "  → Use this if the image is a logo, watermark, company header, signature, "
            "    decorative element, or otherwise unrelated to the question.\n\n"
            "Reply with ONLY the category keyword (TABLE / RELEVANT / IRRELEVANT) "
            "and, for TABLE only, the JSON array below it."
        )
        msg = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=2048,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "image", "source": {
                        "type": "base64", "media_type": media_type, "data": img_b64}},
                    {"type": "text", "text": prompt},
                ],
            }],
        )
        resp = msg.content[0].text.strip()
        if resp.startswith("TABLE"):
            import json
            m = re.search(r'\[[\s\S]*\]', resp)
            if m:
                try:
                    data = json.loads(m.group())
                    if isinstance(data, list) and len(data) > 1:
                        rows = [[str(c) if c is not None else "" for c in row] for row in data]
                        result = ("table", rows)
                    else:
                        result = ("relevant", None)
                except Exception:
                    result = ("relevant", None)
            else:
                result = ("relevant", None)
        elif resp.startswith("RELEVANT"):
            result = ("relevant", None)
        else:
            result = ("irrelevant", None)
    except Exception:
        result = ("irrelevant", None)
    cache[cache_key] = result
    return result


# ── Table quality filter ──────────────────────────────────────────────────────
def _is_quality_table(rows: list) -> bool:
    """
    Return True only for real structured tables.

    Two acceptance paths:
    1. Schedule table: >= 4 standalone x/✓ markers in cells by themselves.
       Protocol text pages never have this many bare 'x' cells.
    2. Bordered/structural table (Method 1): consistent column count across
       rows AND at least 2 multi-word cells (procedure/visit names).
       Word-position noise has wildly varying column counts per row.

    Everything else is rejected as word-position reconstruction noise.
    """
    if not rows or len(rows) < 3:
        return False
    col_counts = [len(r) for r in rows]
    if max(col_counts) < 3:
        return False
    cells = [str(c).strip() for r in rows for c in r if str(c).strip()]
    if not cells:
        return False

    # Path 1 — schedule table: many bare x / ✓ markers
    _MARKERS = {"x", "X", "✓", "•", "x "}
    x_count = sum(1 for c in cells if c in _MARKERS)
    if x_count >= 4:
        return True

    # Reject document-header metadata tables (Protocol No., Version, Date repeated rows)
    meta_rows = sum(
        1 for r in rows
        if r and all(not str(c).strip() or _HEADING_SKIP.search(str(c)) for c in r)
    )
    if meta_rows >= max(1, len(rows) - 1):
        return False

    # Reject page-index / table-of-contents tables
    _TOC_CELL = re.compile(r'^\s*(page\s+\d+|\d+\s*$|p\.\s*\d+)\s*$', re.IGNORECASE)
    toc_cells = sum(1 for c in cells if _TOC_CELL.match(str(c)))
    if toc_cells >= 3:
        return False

    # Path 2 — structural table: consistent column width + multi-word cells
    avg_cols = sum(col_counts) / len(col_counts)
    col_spread = max(col_counts) - min(col_counts)
    multiword = sum(1 for c in cells if " " in c and len(c) > 4)
    if col_spread <= 3 and avg_cols >= 2 and multiword >= 2:
        return True

    # Path 3 — schedule table with numbers as markers (week/visit numbers)
    number_cells = sum(1 for c in cells if re.match(r'^-?\d+$', c))
    multiword2 = sum(1 for c in cells if " " in c)
    if number_cells >= 3 and multiword2 >= 2:
        return True

    return False


# ── Heading helpers ───────────────────────────────────────────────────────────

# Lines that are document-level page headers (repeated on every page in eCRF,
# protocol cover sheets, etc.) — skip when extracting the section heading.
_HEADING_SKIP = re.compile(
    r'protocol\s*(no\.?|number)|document\s+version|e-?\s*crf\s*(completion|guidelines?)?'
    r'|version\s+number|\bdate\s*:|\bsponsor\b|study\s+code|site\s+no'
    r'|completion\s+guide|guidelines?\s*$|form\s+version'
    r'|page\s+\d+\s+of\s+\d+|^\s*\d+\s+of\s+\d+\s*$',
    re.IGNORECASE,
)


def _extract_heading(text: str) -> str:
    """Return the first short line that looks like a section heading,
    skipping repeated document-level header lines (e.g. eCRF page headers)."""
    for line in text.splitlines():
        line = line.strip()
        if (line
                and 4 <= len(line) <= 120
                and not line.endswith('.')
                and not re.match(r'^\d+$', line)
                and not re.match(r'^page\s*\d+$', line, re.I)
                and not _HEADING_SKIP.search(line)):
            return line
    return ""


_H_STOP = {
    "the","a","an","is","in","on","at","to","for","of","and","or","from",
    "with","by","as","be","are","was","this","that","its","not","but",
    "show","give","get","provide","tell","list","what","which","how","me",
}

def _heading_match(query: str, heading: str) -> int:
    """Count overlapping meaningful words (prefix-aware) between query and heading.

    Single-word section headings (e.g. 'Demography', 'Introduction') are treated
    as equivalent to a 2-word match when the query contains that word, because there
    is no way to score higher on a one-word heading.
    """
    def _w(s):
        return {w for w in re.findall(r'\b\w{3,}\b', s.lower()) if w not in _H_STOP}
    q, h = _w(query), _w(heading)
    count = 0
    for qw in q:
        for hw in h:
            if qw == hw or (len(qw) >= 5 and (qw.startswith(hw[:5]) or hw.startswith(qw[:5]))):
                count += 1
                break
    # A single-word heading that matches is as strong as a two-word match
    # (e.g. heading "Demography" matched by query "demography from eCRF guidelines")
    if count >= 1 and len(h) == 1:
        return 2
    return count


def _text_heading_match(query: str, text: str) -> int:
    """Scan the first 15 lines of page text for a section heading that matches
    the query. Used as a fallback when _extract_heading() returned a generic
    document-level title instead of the actual section heading."""
    best = 0
    for line in text.splitlines()[:15]:
        line = line.strip()
        if 3 <= len(line) <= 80 and not line.endswith('.') and not _HEADING_SKIP.search(line):
            score = _heading_match(query, line)
            if score > best:
                best = score
    return best


def get_table_image(filename: str, page_num: int, heading_hint: str = "") -> str | None:
    """
    Render ONLY the relevant table/section from a PDF page, not the whole page.

    Strategy (in order):
    1. find_tables() — if a structured table is detected, crop to its exact bbox.
    2. Text search — search for the heading_hint or "Table" keyword; crop from
       that Y position downward (captures borderless tables).
    3. Full page fallback — if neither works, render the whole page.

    All results are cached by (filename, page_num, hint[:20]).
    """
    cache_key = (filename, page_num, heading_hint[:20])
    cached = st.session_state.page_img_cache.get(cache_key)
    if cached:
        return cached
    pdf_bytes = st.session_state.doc_pdf_bytes.get(filename)
    if not pdf_bytes or not PYMUPDF_OK:
        return None
    try:
        doc  = fitz.open(stream=pdf_bytes, filetype="pdf")
        page = doc[page_num - 1]
        rect = page.rect
        clip = None

        # Strategy 1: structural table bbox via find_tables()
        try:
            tbl_list = page.find_tables()
            if tbl_list.tables:
                # Pick the largest table on the page (most likely the schedule)
                best = max(
                    tbl_list.tables,
                    key=lambda t: (t.bbox[2]-t.bbox[0]) * (t.bbox[3]-t.bbox[1])
                )
                b = best.bbox
                pad = 20
                clip = fitz.Rect(
                    max(0, b[0]-pad), max(0, b[1]-pad),
                    min(rect.width, b[2]+pad), min(rect.height, b[3]+pad)
                )
        except Exception:
            pass

        # Strategy 2: search for the table heading text to find start Y
        if clip is None:
            search_terms = []
            if heading_hint:
                # Use first few meaningful words of the heading
                words = [w for w in heading_hint.split() if len(w) > 3][:4]
                if words:
                    search_terms.append(" ".join(words[:2]))
            search_terms.append("Table")   # generic fallback
            for term in search_terms:
                hits = page.search_for(term)
                if hits:
                    top_y = min(r.y0 for r in hits)
                    clip = fitz.Rect(0, max(0, top_y - 8), rect.width, rect.height)
                    break

        # Render — crop when we have a clip, otherwise full page at lower DPI
        if clip:
            mat = fitz.Matrix(2.2, 2.2)   # 158 DPI — sharp for reading small text
            pix = page.get_pixmap(matrix=mat, clip=clip, alpha=False)
        else:
            mat = fitz.Matrix(1.6, 1.6)
            pix = page.get_pixmap(matrix=mat, alpha=False)

        b64 = base64.b64encode(pix.tobytes("png")).decode()
        doc.close()
        st.session_state.page_img_cache[cache_key] = b64
        return b64
    except Exception:
        return None


def extract_pdf(file_bytes: bytes, filename: str):
    """Extract text, tables, images chunk-by-chunk from PDF."""
    chunks = []
    if not PYMUPDF_OK:
        st.warning("PyMuPDF not installed – install with `pip install pymupdf`")
        return chunks
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    for page_num, page in enumerate(doc, 1):
        text = page.get_text("text").strip()
        images = []
        for img in page.get_images(full=True):
            xref = img[0]
            base = doc.extract_image(xref)
            img_bytes = base["image"]
            images.append(img_to_b64(img_bytes))
        # ── Table extraction: try three methods in order ──────────────────────
        tables = []

        # Method 1: PyMuPDF structural find_tables (best for bordered tables)
        try:
            tbl_list = page.find_tables()
            for t in tbl_list.tables:
                rows = t.extract()
                if rows and len(rows) > 1:
                    tables.append([[str(c) if c is not None else "" for c in row] for row in rows])
        except Exception:
            pass

        # Method 1 is the only extraction method used.
        # Noisy word-position reconstruction (old Methods 2 & 3) has been removed.
        # When Method 1 cannot detect a table (borderless layouts), the page is
        # rendered as a full-fidelity image via get_page_image() at query time,
        # so the user always sees content exactly as it appears in the PDF.
        if text or images or tables:
            chunks.append({
                "text": text,
                "heading": _extract_heading(text),
                "source": filename,
                "page": page_num,
                "images": images,
                "tables": tables,
            })
    return chunks


def extract_docx(file_bytes: bytes, filename: str):
    """Extract text, tables, inline images from DOCX."""
    chunks = []
    if not DOCX_OK:
        st.warning("python-docx not installed – install with `pip install python-docx`")
        return chunks
    from docx.oxml.ns import qn
    doc = DocxDocument(BytesIO(file_bytes))
    text_parts = []
    tables = []
    images = []
    for elem in doc.element.body:
        tag = elem.tag.split("}")[-1]
        if tag == "p":
            para_text = "".join(n.text or "" for n in elem.iter() if n.tag.split("}")[-1] == "t")
            if para_text.strip():
                text_parts.append(para_text)
        elif tag == "tbl":
            rows = []
            for row in elem.iter(qn("w:tr")):
                cells = [
                    "".join(n.text or "" for n in cell.iter() if n.tag.split("}")[-1] == "t")
                    for cell in row.iter(qn("w:tc"))
                ]
                rows.append(cells)
            if rows:
                tables.append(rows)
    # Images
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_bytes = rel.target_part.blob
                images.append(img_to_b64(img_bytes))
            except Exception:
                pass
    body = "\n".join(text_parts)
    chunks.append({
        "text": body,
        "heading": _extract_heading(body),
        "source": filename,
        "page": 1,
        "images": images,
        "tables": tables,
    })
    return chunks


def extract_txt(file_bytes: bytes, filename: str):
    text = file_bytes.decode("utf-8", errors="ignore")
    lines = [l for l in text.split("\n") if l.strip()]
    # split into ~500-char chunks
    chunks = []
    buf = []
    buf_len = 0
    for line in lines:
        buf.append(line)
        buf_len += len(line)
        if buf_len > 800:
            chunks.append({"text": "\n".join(buf), "source": filename, "page": len(chunks)+1, "images": [], "tables": []})
            buf, buf_len = [], 0
    if buf:
        chunks.append({"text": "\n".join(buf), "source": filename, "page": len(chunks)+1, "images": [], "tables": []})
    return chunks


def extract_csv(file_bytes: bytes, filename: str):
    if not PANDAS_OK:
        st.warning("pandas not installed – install with `pip install pandas`")
        return []
    import io
    df = pd.read_csv(io.BytesIO(file_bytes))
    rows = [df.columns.tolist()] + df.values.tolist()
    rows = [[str(c) for c in r] for r in rows]
    return [{"text": df.to_string(index=False), "source": filename, "page": 1, "images": [], "tables": [rows]}]


def extract_excel(file_bytes: bytes, filename: str):
    """Extract each sheet of an Excel file as a separate chunk with a table."""
    if not PANDAS_OK:
        st.warning("pandas not installed – install with `pip install pandas openpyxl`")
        return []
    import io
    chunks = []
    try:
        xl = pd.ExcelFile(io.BytesIO(file_bytes), engine="openpyxl")
    except Exception:
        try:
            xl = pd.ExcelFile(io.BytesIO(file_bytes), engine="xlrd")
        except Exception:
            return []
    for page_num, sheet in enumerate(xl.sheet_names, 1):
        try:
            df = xl.parse(sheet)
            if df.empty:
                continue
            rows = [df.columns.tolist()] + df.values.tolist()
            rows = [[str(c) for c in r] for r in rows]
            chunks.append({
                "text": f"Sheet: {sheet}\n" + df.to_string(index=False),
                "source": filename,
                "page": page_num,
                "images": [],
                "tables": [rows],
            })
        except Exception:
            continue
    return chunks


def load_file(uploaded_file):
    name = uploaded_file.name
    data = uploaded_file.read()
    ext = Path(name).suffix.lower()
    if ext == ".pdf":
        return extract_pdf(data, name)
    elif ext == ".docx":
        return extract_docx(data, name)
    elif ext in (".txt", ".md"):
        return extract_txt(data, name)
    elif ext == ".csv":
        return extract_csv(data, name)
    elif ext in (".xlsx", ".xls"):
        return extract_excel(data, name)
    else:
        return extract_txt(data, name)


@st.cache_resource
def get_embedder():
    if not EMBEDDING_OK:
        return None
    return SentenceTransformer("all-MiniLM-L6-v2")


def embed_chunks(chunks):
    embedder = get_embedder()
    if embedder is None or not chunks:
        return None
    texts = [c["text"][:512] for c in chunks]
    return embedder.encode(texts, show_progress_bar=False)


_DOC_GENERIC = {
    "the", "and", "for", "from", "with", "data", "plan", "review", "file",
    "spec", "guide", "form", "report", "list", "log", "test", "study",
    "clinical", "trial", "version", "final", "draft", "new", "old", "doc",
    "document", "template", "sample", "example", "ref", "information",
}


def _doc_filter(query: str, all_chunks: list) -> tuple:
    """
    Restrict chunks to documents whose name is mentioned in the query.
    Uses a generic-word stoplist so common filename words (data, plan, review…)
    don't cause false matches across unrelated documents.
    Falls back to all chunks only when truly no document name is detected.
    """
    q_lower = query.lower()
    sources = list({c["source"] for c in all_chunks})

    matched = []
    for src in sources:
        stem = Path(src).stem
        # Split on common separators, lowercase each part
        parts = re.split(r'[\s\-_\.]+', stem.lower())
        # Keep only distinctive identifiers: len ≥ 3 and not generic
        distinctive = [p for p in parts if len(p) >= 3 and p not in _DOC_GENERIC]
        # Also test the full stem (space-normalised) as a phrase
        stem_phrase = re.sub(r'[\-_\.]+', ' ', stem.lower())
        identifiers = distinctive + [stem_phrase]
        # Use word-boundary matching so "check" in a filename does not match
        # "checks" in the query (substring match caused Edit Check Spec to be
        # included when user asked "from the DRP" with "checks" in the query).
        if any(
            ident and re.search(r'\b' + re.escape(ident) + r'\b', q_lower)
            for ident in identifiers
        ):
            matched.append(src)

    if not matched:
        # For schedule/table queries prefer protocol-named documents
        SCHED_Q = re.compile(
            r'\b(assessment\s+schedule|schedule\s+of\s+assessments?|'
            r'visit\s+schedule|time\s+and\s+events?|table\s+\d)\b',
            re.IGNORECASE
        )
        if SCHED_Q.search(query):
            proto_docs = [s for s in sources
                          if re.search(r'protocol', Path(s).stem, re.IGNORECASE)]
            if proto_docs:
                matched = proto_docs

    matched_set = set(matched) if matched else {c["source"] for c in all_chunks}
    filtered = [(i, c) for i, c in enumerate(all_chunks) if c["source"] in matched_set]
    indices = [i for i, _ in filtered]
    pool    = [c for _, c in filtered]
    return pool, indices


# ── CDISC domain filter ───────────────────────────────────────────────────────
# When the query explicitly names a domain (e.g. "DM domain", "VS domain"),
# keep only chunks that specifically discuss that domain.
_CDISC_DOMAINS = frozenset({
    "DM", "AE", "VS", "LB", "CM", "EX", "DS", "MH", "SU", "EG", "PE", "SC",
    "RP", "FA", "PC", "PP", "DV", "QS", "PR", "DD", "IE", "SV", "TU", "TR",
    "RS", "MI", "NV", "OE", "HR", "CE", "ML", "BE", "BW",
})


def _detect_domain(query: str) -> str | None:
    """Return the first CDISC domain code found as a whole word in the query."""
    q_upper = query.upper()
    for d in _CDISC_DOMAINS:
        if re.search(r'\b' + re.escape(d) + r'\b', q_upper):
            return d
    return None


_DOMAIN_COL_RE = re.compile(r'^\s*domain\s*$', re.IGNORECASE)

def _filter_table_rows_by_domain(rows: list, domain: str) -> list:
    """
    Keep only the header row plus data rows whose Domain or Dataset column value
    matches the requested domain (or 'ALL').

    Checks two column types:
    - 'Domain' column: used by Edit Check Spec tables (values like "DM", "AE")
    - 'Dataset*' column: used by DRP tables where Domain col has descriptive names
      ("Vitals") but the CDISC code ("VS") lives in Dataset / CRF Page column.

    Falls back to whole-row search only when neither column type is found.
    Returns [header] only (no data rows) when nothing matches — caller skips
    via len < 2.
    """
    if not rows or len(rows) < 2:
        return rows

    header = rows[0]
    dom_re = re.compile(r'\b' + re.escape(domain) + r'\b', re.IGNORECASE)
    all_re = re.compile(r'^\s*all\s*$', re.IGNORECASE)

    def _cell_matches(val):
        s = str(val)
        return dom_re.search(s) or all_re.match(s)

    domain_col = next(
        (i for i, cell in enumerate(header) if _DOMAIN_COL_RE.match(str(cell))),
        None,
    )
    # DRP-style tables: CDISC code lives in "Dataset / CRF Page" column
    dataset_col = next(
        (i for i, cell in enumerate(header)
         if re.match(r'^\s*dataset\b', str(cell), re.IGNORECASE)),
        None,
    )

    if domain_col is not None or dataset_col is not None:
        matched = [
            r for r in rows[1:]
            if (
                (domain_col is not None and domain_col < len(r) and _cell_matches(r[domain_col]))
                or
                (dataset_col is not None and dataset_col < len(r) and _cell_matches(r[dataset_col]))
            )
        ]
    else:
        # No explicit domain/dataset column — fall back to whole-row search
        matched = [r for r in rows[1:] if any(dom_re.search(str(cell)) for cell in r)]

    return [header] + matched


def _domain_filter(query: str, chunks: list):
    """
    If the query names a CDISC domain, return only chunks containing that domain.
    Returns None when no domain is detected (caller keeps original list).
    """
    q_upper = query.upper()
    matched_domain = None
    for d in _CDISC_DOMAINS:
        if re.search(r'\b' + re.escape(d) + r'\b', q_upper):
            matched_domain = d
            break
    if not matched_domain:
        return None
    dom_re = re.compile(r'\b' + re.escape(matched_domain) + r'\b')
    hits = [
        c for c in chunks
        if dom_re.search(c.get("text", "")) or dom_re.search(c.get("heading", ""))
    ]
    return hits if hits else None


def _expand_table_refs(retrieved: list, pool: list) -> list:
    """
    Expand retrieved chunks with adjacent pages when table-related content is referenced.
    Looks both forward and backward to handle tables that appear before/after references.
    """
    import re

    TABLE_REF   = re.compile(
        r'\bTable\s+[\d]+(?:[.\-–][\d]+)?|\bFigure\s+[\d]+(?:[.\-–][\d]+)?',
        re.IGNORECASE
    )
    SCHEDULE_KW = re.compile(
        r'\b(schedule\s+of\s+assessments?|assessment\s+schedule|'
        r'visit\s+schedule|study\s+schedule|time\s+and\s+events?|'
        r'procedures?\s+and\s+assessments?)\b',
        re.IGNORECASE
    )

    by_src_page = {(c["source"], c["page"]): c for c in pool}
    seen  = {(c["source"], c["page"]) for c in retrieved}
    extra = []
    _MAX_EXTRA = 10  # never add more than 10 pages via expansion

    def _add_window(src, center_page, before, after):
        for offset in range(-before, after + 1):
            if len(extra) >= _MAX_EXTRA:
                return
            if offset == 0:
                continue
            key = (src, center_page + offset)
            if key not in by_src_page or key in seen:
                continue
            extra.append(by_src_page[key])
            seen.add(key)

    for chunk in retrieved:
        src, page = chunk["source"], chunk["page"]
        text = chunk["text"]

        # Explicit table reference → look ±pages (table may be before or after the text)
        # Note: always expand even if the current chunk already has some table data
        if TABLE_REF.search(text):
            _add_window(src, page, before=2, after=8)

        # Schedule section heading: table is usually 1-5 pages after
        if SCHEDULE_KW.search(text):
            _add_window(src, page, before=1, after=6)

    # Scan entire pool for x-grid pages (assessment schedule columns of x markers)
    X_GRID = re.compile(r'(?:^|\s)[xX](?:[\s\t]+[xX]){2,}', re.MULTILINE)
    for c in pool:
        key = (c["source"], c["page"])
        if key not in seen and X_GRID.search(c.get("text", "")):
            extra.append(c)
            seen.add(key)

    return retrieved + extra


_STOP_WORDS = {
    "the","a","an","is","in","on","at","to","for","of","and","or","from",
    "with","by","as","be","are","was","were","this","that","it","its",
    "have","has","not","but","what","which","who","when","where","how",
    "give","show","me","get","find","provide","tell","list","all","only",
}

def retrieve(query: str, chunks, embeddings, top_k=5):
    if not chunks:
        return []

    # Narrow to the document named in the query (e.g. "from protocol", "from DRP")
    pool, pool_idx = _doc_filter(query, chunks)

    # Meaningful query words only (strip stop words + short tokens)
    q_words = {w for w in query.lower().split()
               if w not in _STOP_WORDS and len(w) > 2}
    if not q_words:                          # fallback if all words stripped
        q_words = set(query.lower().split())

    # Keyword pre-score on the pool
    kw_scores = []
    for local_i, c in enumerate(pool):
        t = c["text"].lower()
        score = sum(t.count(w) for w in q_words)
        kw_scores.append((score, local_i))
    kw_scores.sort(key=lambda x: -x[0])

    if embeddings is None or not EMBEDDING_OK:
        result = [pool[i] for s, i in kw_scores[:top_k] if s > 0]
        return result or pool[:top_k]

    import numpy as np

    # Candidate selection:
    # - If pool is small (single doc after filter) → search ALL chunks semantically
    # - If pool is large (all docs) → keyword pre-filter top-50 then semantic
    if len(pool) <= 60:
        candidate_local = list(range(len(pool)))
    else:
        candidate_local = [i for _, i in kw_scores[:50]] or list(range(min(50, len(pool))))

    candidate_global = [pool_idx[i] for i in candidate_local]

    # Cache query embedding
    if "q_emb_cache" not in st.session_state:
        st.session_state.q_emb_cache = {}
    if query not in st.session_state.q_emb_cache:
        embedder = get_embedder()
        st.session_state.q_emb_cache[query] = embedder.encode(
            [query], show_progress_bar=False, convert_to_numpy=True
        )
    q_emb = st.session_state.q_emb_cache[query]

    candidate_embs = embeddings[candidate_global]
    scores = np.dot(candidate_embs, q_emb.T).flatten()
    top_local = scores.argsort()[::-1][:top_k]

    result = [pool[candidate_local[i]] for i in top_local]
    result = result or pool[:top_k]

    # Heading-based expansion: find pages whose section heading directly matches
    # the query (e.g. "9.4.2 Statistical model, hypothesis, and method of analysis").
    # Semantic search may miss the actual section page if it's outranked by pages
    # that scatter the same keywords throughout body text.
    # Uses both _heading_match (on extracted heading) and _text_heading_match
    # (scans first 15 lines) so pages with repeated document headers (eCRF etc.)
    # are still found even if _extract_heading returned the wrong line.
    _seen_hm  = {(c["source"], c["page"]) for c in result}
    _by_sp    = {(c["source"], c["page"]): c for c in pool}

    def _best_hmatch(c):
        return max(
            _heading_match(query, c.get("heading", "")),
            _text_heading_match(query, c.get("text", "")),
        )

    _hm_hits = sorted(
        [(c, _best_hmatch(c)) for c in pool if _best_hmatch(c) >= 2],
        key=lambda x: -x[1]
    )
    for _hm_c, _ in _hm_hits[:3]:
        sp = (_hm_c["source"], _hm_c["page"])
        if sp not in _seen_hm:
            result.append(_hm_c)
            _seen_hm.add(sp)
        # Add the following page in case the section continues
        sp1 = (_hm_c["source"], _hm_c["page"] + 1)
        if sp1 in _by_sp and sp1 not in _seen_hm:
            result.append(_by_sp[sp1])
            _seen_hm.add(sp1)

    # Schedule queries are handled by the targeted anchor block below —
    # skip broad _expand_table_refs to avoid pulling in unrelated pages.
    _TABLE_Q_INNER = re.compile(
        r'\b(assessment\s+schedule|schedule\s+of\s+assessments?|'
        r'visit\s+schedule|time\s+and\s+events?|'
        r'procedures?\s+and\s+assessments?)\b',
        re.IGNORECASE
    )

    # ── Targeted schedule/table search ──────────────────────────────────────
    # Only triggered when the query explicitly asks for a schedule or table.
    # KEY CHANGE: anchor uses schedule-specific keywords ONLY — NOT generic
    # "Table X.X" references (those appear on every page of a protocol and
    # previously caused the entire document to be returned).
    TABLE_Q = re.compile(
        r'\b(assessment\s+schedule|schedule\s+of\s+assessments?|'
        r'visit\s+schedule|time\s+and\s+events?|'
        r'procedures?\s+and\s+assessments?)\b',
        re.IGNORECASE
    )
    if TABLE_Q.search(query):
        by_src_page = {(c["source"], c["page"]): c for c in pool}
        SCHED_HEADING = re.compile(
            r'\b(schedule\s+of\s+assessments?|assessment\s+schedule|'
            r'time\s+and\s+events?\s+table|visit\s+schedule)\b',
            re.IGNORECASE
        )
        X_GRID = re.compile(r'(?:^|\s)[xX](?:[\s\t]+[xX]){2,}', re.MULTILINE)
        TOC_RE = re.compile(r'\.{3,}\s*\d+', re.MULTILINE)

        def _anchor_score(sp):
            c = by_src_page.get(sp)
            if not c:
                return 0
            return (2 if c.get("tables") else 0) + (1 if X_GRID.search(c.get("text", "")) else 0)

        # All pages that mention the schedule heading
        heading_pages = sorted(
            {(c["source"], c["page"]) for c in pool
             if SCHED_HEADING.search(c.get("text", ""))},
            key=lambda x: x[1]
        )

        # Split: pages that HAVE actual table data vs. TOC/mention-only pages
        real_pages = [(s, p) for s, p in heading_pages if _anchor_score((s, p)) > 0]
        toc_pages  = [(s, p) for s, p in heading_pages if _anchor_score((s, p)) == 0]

        sched_result: list = []
        seen_sched: set = set()

        if real_pages:
            # Expand slightly from each real table page (handles multi-page tables)
            for src, page in real_pages[:2]:
                for offset in range(-1, 4):
                    key = (src, page + offset)
                    if key in by_src_page and key not in seen_sched:
                        sched_result.append(by_src_page[key])
                        seen_sched.add(key)
        else:
            # Only TOC mentions found — search forward from them, keep only table pages
            for src, page in toc_pages[:2]:
                for offset in range(0, 16):
                    key = (src, page + offset)
                    if key in by_src_page and key not in seen_sched and _anchor_score(key) > 0:
                        sched_result.append(by_src_page[key])
                        seen_sched.add(key)

        if sched_result:
            # Return ONLY the schedule table pages (max 4 for multi-page tables)
            result = sorted(
                sched_result,
                key=lambda c: (-_anchor_score((c["source"], c["page"])), c["page"])
            )[:4]
            return result  # early return — skip broad final trim

    # ── Final trim: never send more than 12 chunks to the LLM ───────────────
    # Prioritise: pages with tables > pages matching schedule keywords > rest
    if len(result) > 12:
        SCHED_KW = re.compile(
            r'\b(schedule|assessment|visit|procedure|endpoint)\b', re.IGNORECASE
        )
        def _priority(c):
            has_table   = 1 if c.get("tables") else 0
            has_kw      = 1 if SCHED_KW.search(c.get("text", "")) else 0
            return (has_table * 2 + has_kw, c["page"])

        if TABLE_Q.search(query):
            result = sorted(result, key=_priority, reverse=True)[:12]
        else:
            result = result[:12]

    return result


def build_context(relevant_chunks):
    parts = []
    for c in relevant_chunks:
        part = f"[Source: {c['source']}, Page: {c['page']}]\n{c['text']}"
        if c.get("tables"):
            for t in c["tables"]:
                if t and len(t) >= 1:
                    # Proper markdown table: header | sep | rows
                    cols = len(t[0])
                    header  = "| " + " | ".join(str(x) for x in t[0]) + " |"
                    sep     = "| " + " | ".join("---" for _ in range(cols)) + " |"
                    rows    = "\n".join(
                        "| " + " | ".join(str(x) for x in r) + " |"
                        for r in t[1:]
                    )
                    part += f"\n\n[TABLE]\n{header}\n{sep}\n{rows}\n[/TABLE]"
        parts.append(part)
    return "\n\n---\n\n".join(parts)


def ask_llm(query: str, context: str, vision_images: list = None, table_query: bool = False) -> str:
    """
    vision_images: optional list of (b64, src, page) for relevant non-table images.
    table_query:   True when the user is asking for a schedule/table — triggers a
                   strict "table only, no surrounding text" system prompt.
    """
    if not ANTHROPIC_OK:
        return f"(Anthropic SDK not installed)\n\nRelevant content found:\n\n{context[:2000]}"
    api_key = st.session_state.get("anthropic_api_key", "") or os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return "⚠️ Please enter your Anthropic API key in the sidebar to enable AI answers."
    client = anthropic.Anthropic(api_key=api_key)

    if table_query:
        system = (
            "You are a precise table extraction assistant.\n\n"
            "DOCUMENT RULES:\n"
            "- Each context block is labelled [Source: filename, Page: N].\n"
            "- If the user names a specific document, use ONLY chunks from that document.\n"
            "- NEVER mix data from different documents.\n\n"
            "YOUR ONLY JOB FOR THIS REQUEST:\n"
            f"- The user asked: \"{query}\"\n"
            "- Identify the ONE table whose title or heading best matches the user's request.\n"
            "- DO NOT return tables whose headings do not match — e.g. if the user asks for the "
            "  'assessment schedule', do NOT return eligibility criteria tables, endpoint tables, "
            "  or any other table.\n"
            "- Reproduce that matched table as a COMPLETE markdown table — every column header, "
            "  every row, every cell. Do not skip any rows.\n"
            "- If the matching table spans multiple context blocks, merge them into one table.\n"
            "- If an image is provided and its heading matches the request, extract it as markdown.\n\n"
            "OUTPUT FORMAT — strictly:\n"
            "**[Exact table title from document]**\n"
            "<complete markdown table>\n"
            "[Source: filename, Page: N]\n\n"
            "FORBIDDEN:\n"
            "- Do NOT include any introductory sentences, explanations, footnotes, or surrounding text.\n"
            "- Do NOT return unrelated tables.\n"
            "- Do NOT summarise or abbreviate any rows.\n"
            "- Do NOT output anything except the heading, the markdown table, and the source citation."
        )
        instruction = (
            f"Return ONLY the table whose heading matches: \"{query}\". "
            "No other tables. No surrounding text."
        )
    else:
        # ── Detect CDISC domain in query for hard domain filter ───────────────
        _detected_domain = _detect_domain(query)

        _domain_rule = ""
        _domain_instr = ""
        if _detected_domain:
            _other_domains = ", ".join(
                d for d in ["DM","AE","VS","LB","CM","EX","DS","MH","SU","EG","PE","SC"]
                if d != _detected_domain
            )
            _domain_rule = (
                f"\nDOMAIN FILTER (critical):\n"
                f"- The user is asking specifically about the **{_detected_domain}** domain.\n"
                f"- Output ONLY the edit checks / rules / fields / information whose "
                f"  Domain column value is '{_detected_domain}' or 'ALL'.\n"
                f"- COMPLETELY IGNORE rows whose Domain column is any other value "
                f"  ({_other_domains}, etc.) even if they mention {_detected_domain} "
                f"  in a description or cross-reference field.\n"
                f"- If no {_detected_domain} domain rows are found, say: "
                f"  'No {_detected_domain} domain information found in the document.'\n"
            )
            _domain_instr = (
                f" Output ONLY rows where Domain='{_detected_domain}' or Domain='ALL' — "
                f"skip every other domain row completely."
            )

        system = (
            "You are a precise document extraction assistant. Return exact content — never paraphrase.\n\n"
            "DOCUMENT RULES (most important):\n"
            "- Each context block is labelled [Source: filename, Page: N].\n"
            "- If the user names a specific document (e.g. 'from protocol', 'from DRP'), use ONLY "
            "  chunks whose Source matches that document. Ignore all other sources completely.\n"
            "- If the requested information is not found in the specified document, say exactly: "
            "  'This information is not available in [document name].'\n"
            "- NEVER pull information from a different document to fill gaps.\n\n"
            + _domain_rule +
            "\nTOPIC FOCUS:\n"
            "- Answer ONLY the specific topic asked. Do NOT reproduce surrounding paragraphs, "
            "  unrelated sections, or the full document text.\n"
            "- Extract the minimum content that directly answers the question.\n"
            "- If the user names multiple documents, extract the topic from each separately.\n\n"
            "ANSWER LENGTH RULES:\n"
            "- For factual questions (dosage, date, name, number, eligibility criterion, etc.): "
            "  answer in 1-5 sentences. Quote the exact text from the document, nothing more.\n"
            "- NEVER dump entire sections or chapters. If the relevant answer is one sentence, "
            "  return one sentence — not the surrounding page.\n\n"
            "TABLE RULES:\n"
            "- Every [TABLE]...[/TABLE] block → reproduce as a complete markdown table, every row.\n"
            "- Columnar plain text → reconstruct as markdown table.\n"
            "- If images are provided, read them; if they contain tables extract as markdown table.\n\n"
            "FORMAT:\n"
            "1. Bold heading matching the topic.\n"
            "2. Exact content (concise for facts, complete for tables).\n"
            "3. [Source: filename, Page: N]"
        )
        instruction = (
            "Extract ONLY the content about the specific topic asked, "
            "from the correct source document only. "
            "Reproduce tables in full as markdown tables."
            + _domain_instr
        )

    # Build message content — include vision images so Claude can read them
    user_content: list = []
    if vision_images:
        for b64, src, page in vision_images[:6]:
            user_content.append({
                "type": "image",
                "source": {"type": "base64", "media_type": _get_media_type(b64), "data": b64},
            })
    user_content.append({
        "type": "text",
        "text": (
            f"Documents:\n{context}\n\n"
            f"User question: {query}\n\n"
            f"Instruction: {instruction}"
        ),
    })
    try:
        msg = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=4096,
            system=system,
            messages=[{"role": "user", "content": user_content}],
        )
        return msg.content[0].text
    except Exception as _e:
        _em = str(_e).lower()
        if "401" in _em or "authentication" in _em or "api_key" in _em or "x-api-key" in _em:
            return "⚠️ **Invalid API key.** Please enter a valid Anthropic API key in the sidebar."
        if "429" in _em or "rate_limit" in _em:
            return "⚠️ **Rate limit reached.** Please wait a moment and try again."
        if "overloaded" in _em or "529" in _em:
            return "⚠️ **Anthropic API is temporarily overloaded.** Please retry in a few seconds."
        return f"⚠️ **API error:** {str(_e)}"


def table_to_html(rows):
    if not rows:
        return ""
    html = "<table>"
    for i, row in enumerate(rows):
        html += "<tr>"
        tag = "th" if i == 0 else "td"
        for cell in row:
            html += f"<{tag}>{cell}</{tag}>"
        html += "</tr>"
    html += "</table>"
    return html




def render_answer(answer: str, relevant_chunks):
    """Strip markdown tables from bubble text; collect ALL tables as dataframes."""
    lines = answer.split("\n")
    non_table_lines: list = []
    _tbl_buf: list = []        # accumulates rows of the current markdown table
    md_tables: list = []       # parsed tables from LLM markdown output

    for line in lines:
        stripped = line.strip()
        is_data_row = bool(re.match(r"^\|.+\|$", stripped)) and not re.match(r"^\|[\s\-|:]+\|$", stripped)
        is_sep_row  = bool(re.match(r"^\|[\s\-|:]+\|$", stripped))
        if is_data_row:
            _tbl_buf.append(stripped)
        elif is_sep_row:
            pass  # skip separator, keep buffer open
        else:
            if _tbl_buf:
                rows = [[c.strip() for c in tl.strip("|").split("|")] for tl in _tbl_buf]
                if len(rows) >= 2:
                    md_tables.append(rows)
                _tbl_buf = []
            non_table_lines.append(line)
    if _tbl_buf:
        rows = [[c.strip() for c in tl.strip("|").split("|")] for tl in _tbl_buf]
        if len(rows) >= 2:
            md_tables.append(rows)

    # Collapse consecutive blank lines
    cleaned: list[str] = []
    prev_blank = False
    for line in non_table_lines:
        is_blank = line.strip() == ""
        if is_blank and prev_blank:
            continue
        cleaned.append(line)
        prev_blank = is_blank
    answer_html = "<br>".join(cleaned).strip()

    # Images — deduplicate; logos appear on every page so skip them
    from collections import Counter
    _img_count: Counter = Counter()
    _img_first: dict = {}
    for c in relevant_chunks:
        for img_b64 in c.get("images", []):
            _img_count[img_b64] += 1
            if img_b64 not in _img_first:
                _img_first[img_b64] = (img_b64, c["source"], c["page"])
    images = [info for b64, info in _img_first.items() if _img_count[b64] == 1]

    # Tables: chunk-extracted (structured) first, then LLM-generated markdown tables
    seen = set()
    extra_tables = []
    for c in relevant_chunks:
        for t in c.get("tables", []):
            key = (c["source"], c["page"])
            if key not in seen:
                seen.add(key)
                extra_tables.append((t, c["source"], c["page"]))
    # Add LLM markdown tables that have no corresponding chunk-extracted table
    # (this happens when find_tables() failed but LLM reconstructed from plain text)
    for rows in md_tables:
        extra_tables.append((rows, "protocol", 0))

    return answer_html, images, extra_tables


def _md_table_to_html(lines):
    if not lines:
        return ""
    # Filter separator lines (---|---|---)
    rows = [l for l in lines if not re.match(r"^\|[\s\-|:]+\|$", l)]
    html = "<table>"
    for i, row in enumerate(rows):
        cells = [c.strip() for c in row.strip("|").split("|")]
        html += "<tr>"
        tag = "th" if i == 0 else "td"
        for c in cells:
            html += f"<{tag}>{c}</{tag}>"
        html += "</tr>"
    html += "</table>"
    return html


# ═══════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ═══════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("""
    <div class="sidebar-brand">
        <h1>📚 StudyGuide</h1>
        <p>Document Intelligence Assistant</p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("#### Upload Documents")
    uploaded_files = st.file_uploader(
        "Drop files here",
        type=["pdf", "docx", "txt", "md", "csv", "xlsx", "xls"],
        accept_multiple_files=True,
        label_visibility="collapsed",
    )

    # Folder path only works when running locally (not on Streamlit Cloud)
    _is_cloud = os.environ.get("STREAMLIT_SERVER_HEADLESS", "").lower() in ("1", "true")
    if _is_cloud:
        folder_path = ""
        st.markdown(
            '<div style="font-size:.75rem;color:var(--text-muted);padding:6px 10px;'
            'background:var(--surface2);border:1px solid var(--border);border-radius:8px;'
            'margin-top:4px;">📁 Folder path loading is only available when running locally.</div>',
            unsafe_allow_html=True,
        )
    else:
        _default_folder = st.secrets.get("DOCS_FOLDER", r"C:\Users\User\OneDrive\Desktop\Study Guide\docs")
        folder_path = st.text_input(
            "Or enter a folder path",
            placeholder="C:/path/to/your/documents",
            value=_default_folder,
        )

    col1, col2 = st.columns(2)
    with col1:
        load_btn = st.button("⚡ Load", use_container_width=True)
    with col2:
        clear_btn = st.button("🗑 Clear", use_container_width=True)

    if clear_btn:
        st.session_state.doc_chunks = []
        st.session_state.embeddings = None
        st.session_state.loaded_files = []
        st.session_state.messages = []
        st.session_state.doc_pdf_bytes = {}
        st.session_state.page_img_cache = {}
        st.session_state.learned_answers = []
        st.session_state.pop("_awaiting_feedback", None)
        st.session_state.pop("_feedback_query", None)
        st.session_state.pop("_retry_mode", None)
        st.rerun()

    if load_btn:
        import concurrent.futures

        EXTS = {".pdf", ".docx", ".txt", ".md", ".csv", ".xlsx", ".xls"}

        # Collect all files to process (uploaded + folder), skip already-loaded
        tasks = []  # list of (name, bytes)

        if uploaded_files:
            for f in uploaded_files:
                if f.name not in st.session_state.loaded_files:
                    data = f.read()
                    tasks.append((f.name, data))
                    # Store raw PDF bytes so pages can be rendered as images later
                    if f.name.lower().endswith(".pdf"):
                        st.session_state.doc_pdf_bytes[f.name] = data

        if folder_path and Path(folder_path).is_dir():
            for fp in Path(folder_path).rglob("*"):
                if fp.suffix.lower() in EXTS and fp.name not in st.session_state.loaded_files:
                    try:
                        data = fp.read_bytes()
                        tasks.append((fp.name, data))
                        if fp.suffix.lower() == ".pdf":
                            st.session_state.doc_pdf_bytes[fp.name] = data
                    except Exception:
                        pass

        if not tasks and not uploaded_files and not folder_path:
            st.warning("Please upload files or enter a folder path.")
        elif tasks:
            def _parse(name_bytes):
                name, data = name_bytes
                class _F:
                    def __init__(self): self.name = name
                    def read(self): return data
                try:
                    return name, load_file(_F())
                except Exception:
                    return name, []

            new_chunks = []
            with st.spinner(f"Reading {len(tasks)} file(s)…"):
                # Parallel parse — I/O already done above, this parallelises CPU parsing
                with concurrent.futures.ThreadPoolExecutor(max_workers=min(8, len(tasks))) as ex:
                    for name, chunks in ex.map(_parse, tasks):
                        new_chunks.extend(chunks)
                        st.session_state.loaded_files.append(name)

            if new_chunks:
                st.session_state.doc_chunks.extend(new_chunks)
                # Only embed the NEW chunks, then concatenate with existing embeddings
                with st.spinner(f"Indexing {len(new_chunks)} chunks…"):
                    new_embs = embed_chunks(new_chunks)
                    if new_embs is not None:
                        import numpy as np
                        if st.session_state.embeddings is not None:
                            st.session_state.embeddings = np.vstack(
                                [st.session_state.embeddings, new_embs]
                            )
                        else:
                            st.session_state.embeddings = new_embs
                st.success(f"Loaded {len(new_chunks)} chunks from {len(tasks)} file(s)!")

    st.markdown("---")

    # Stats
    n_files = len(set(c["source"] for c in st.session_state.doc_chunks)) if st.session_state.doc_chunks else 0
    n_chunks = len(st.session_state.doc_chunks)
    n_imgs = sum(len(c["images"]) for c in st.session_state.doc_chunks)

    c1, c2 = st.columns(2)
    with c1:
        st.markdown(f'<div class="stat-card"><div class="num">{n_files}</div><div class="lbl">Files</div></div>', unsafe_allow_html=True)
    with c2:
        st.markdown(f'<div class="stat-card"><div class="num">{n_chunks}</div><div class="lbl">Chunks</div></div>', unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown(f'<div class="stat-card"><div class="num">{n_imgs}</div><div class="lbl">Images Indexed</div></div>', unsafe_allow_html=True)

    if st.session_state.loaded_files:
        st.markdown("---")
        st.markdown(
            '<p style="font-size:.85rem;font-weight:700;color:#f0eeff;'
            'margin:0 0 8px;letter-spacing:.04em;text-transform:uppercase;">Loaded Files</p>',
            unsafe_allow_html=True,
        )
        icons = {"pdf":"📄","docx":"📝","txt":"📃","md":"📋","csv":"📊","xlsx":"📗","xls":"📗"}
        for fname in st.session_state.loaded_files:
            ext = Path(fname).suffix.lower().lstrip(".")
            icon = icons.get(ext, "📁")
            st.markdown(
                f'<div style="display:flex;align-items:center;gap:6px;'
                f'padding:5px 8px;margin-bottom:4px;border-radius:6px;'
                f'background:rgba(255,255,255,.05);border:1px solid rgba(255,255,255,.08);">'
                f'<span style="font-size:.9rem;">{icon}</span>'
                f'<span style="font-size:.78rem;color:#c8c5ff;word-break:break-all;">{fname}</span>'
                f'</div>',
                unsafe_allow_html=True,
            )

    st.markdown("---")
    st.markdown('<span style="font-size:.75rem;color:var(--text-muted);letter-spacing:.04em;text-transform:uppercase;">API Key</span>', unsafe_allow_html=True)
    # Load API key: sidebar input > session state > env var > Streamlit secrets
    _default_key = (
        st.session_state.get("anthropic_api_key", "")
        or os.environ.get("ANTHROPIC_API_KEY", "")
        or st.secrets.get("ANTHROPIC_API_KEY", "")
    )
    api_key_input = st.text_input(
        "Anthropic API Key",
        type="password",
        placeholder="sk-ant-…",
        label_visibility="collapsed",
        value=_default_key,
    )
    if api_key_input:
        st.session_state["anthropic_api_key"] = api_key_input

# ═══════════════════════════════════════════════════════════════════════════════
# MAIN AREA
# ═══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<div class="main-header">
    <h2>Ask Your Documents</h2>
    <p>Exact answers · Tables preserved · Images displayed</p>
</div>
""", unsafe_allow_html=True)

# Welcome state
if not st.session_state.doc_chunks:
    st.markdown("""
    <div style="text-align:center; padding:4rem 2rem; color:var(--text-muted);">
        <div style="font-size:4rem; margin-bottom:1rem;">📂</div>
        <div style="font-size:1.1rem; font-family:'Playfair Display',serif; color:var(--text); margin-bottom:.5rem;">
            No documents loaded yet
        </div>
        <div style="font-size:.85rem;">Upload files in the sidebar to get started.</div>
        <br>
        <div style="display:flex; justify-content:center; gap:1rem; flex-wrap:wrap; margin-top:1rem;">
            <span style="background:var(--surface); border:1px solid var(--border); border-radius:20px; padding:6px 16px; font-size:.78rem;">📄 PDF</span>
            <span style="background:var(--surface); border:1px solid var(--border); border-radius:20px; padding:6px 16px; font-size:.78rem;">📝 DOCX</span>
            <span style="background:var(--surface); border:1px solid var(--border); border-radius:20px; padding:6px 16px; font-size:.78rem;">📃 TXT / MD</span>
            <span style="background:var(--surface); border:1px solid var(--border); border-radius:20px; padding:6px 16px; font-size:.78rem;">📊 CSV</span>
            <span style="background:var(--surface); border:1px solid var(--border); border-radius:20px; padding:6px 16px; font-size:.78rem;">📗 XLSX / XLS</span>
        </div>
    </div>
    """, unsafe_allow_html=True)
else:
    # Chat history
    st.markdown('<div class="chat-wrap">', unsafe_allow_html=True)
    for _msg_idx, msg in enumerate(st.session_state.messages):
        _is_last = _msg_idx == len(st.session_state.messages) - 1
        if msg["role"] == "user":
            _q_text = msg["content"]
            # Bubble + action buttons on the same row (bubble wide, buttons narrow right)
            _qbubble_col, _qcopy_col, _qreload_col = st.columns([9.0, 0.5, 0.5])
            with _qbubble_col:
                st.markdown(
                    f'<div class="msg-user"><div class="bubble">{_q_text}</div></div>',
                    unsafe_allow_html=True,
                )
            with _qcopy_col:
                st.markdown(
                    '<div style="display:flex;align-items:center;justify-content:center;height:100%;padding-top:6px;">',
                    unsafe_allow_html=True,
                )
                if st.button("📋", key=f"copy_q_{_msg_idx}",
                             use_container_width=True, help="Copy question"):
                    _ck = f"_show_copy_q_{_msg_idx}"
                    st.session_state[_ck] = not st.session_state.get(_ck, False)
                st.markdown('</div>', unsafe_allow_html=True)
            with _qreload_col:
                st.markdown(
                    '<div style="display:flex;align-items:center;justify-content:center;height:100%;padding-top:6px;">',
                    unsafe_allow_html=True,
                )
                if st.button("🔄", key=f"reload_q_{_msg_idx}",
                             use_container_width=True, help="Re-ask this question"):
                    st.session_state["_pending_query"] = _q_text
                    st.rerun()
                st.markdown('</div>', unsafe_allow_html=True)
            if st.session_state.get(f"_show_copy_q_{_msg_idx}"):
                st.caption("Your question — click the copy icon to copy:")
                st.code(_q_text, language=None)
        else:
            answer_html = msg.get("answer_html", msg["content"])
            images = msg.get("images", [])
            tables = msg.get("tables", [])

            st.markdown(f"""
            <div class="msg-bot">
                <div class="avatar">🤖</div>
                <div class="bubble" id="ans-bubble-{_msg_idx}">
                    {answer_html}
                </div>
            </div>
            """, unsafe_allow_html=True)

            # Copy answer — st.expander shows the text with Streamlit's built-in copy icon
            with st.expander("📋 Copy Answer Text"):
                st.caption("Answer text — click the copy icon to copy:")
                st.code(msg.get("content", ""), language=None)

            # Inline images (page renders + classified doc images)
            if images:
                for b64, src, page in images:
                    st.image(
                        base64.b64decode(b64),
                        caption=f"{src} – p.{page}",
                        use_container_width=True,
                    )

            # Structured tables — plain Streamlit dataframe
            for (t, src, page) in tables:
                if t and len(t) > 1:
                    try:
                        import pandas as pd
                        df = pd.DataFrame(t[1:], columns=[str(c) for c in t[0]])
                        st.dataframe(df, use_container_width=True)
                    except Exception:
                        st.markdown(table_to_html(t), unsafe_allow_html=True)

            # ── Feedback buttons (last assistant message only) ────────────────
            if _is_last and st.session_state.get("_awaiting_feedback"):
                st.markdown("""<style>
/* Feedback buttons — white background, black text, clearly visible on dark bg */
div[data-testid="stHorizontalBlock"] div[data-testid="stButton"] button {
    background: #ffffff !important;
    border: 1.5px solid #cccccc !important;
    color: #111111 !important;
    -webkit-text-fill-color: #111111 !important;
    font-weight: 600 !important;
    font-size: 0.92rem !important;
    border-radius: 8px !important;
    min-height: 44px !important;
    height: 44px !important;
    padding: 8px 20px !important;
    line-height: 1.2 !important;
    width: 100% !important;
}
div[data-testid="stHorizontalBlock"] div[data-testid="stButton"] button:hover {
    background: #f0f0f0 !important;
    border-color: #888888 !important;
    color: #000000 !important;
    -webkit-text-fill-color: #000000 !important;
}
</style>""", unsafe_allow_html=True)
                st.markdown(
                    '<div style="font-size:.88rem;color:#f0eeff;font-weight:700;'
                    'margin:8px 0 6px 0;">Is this information correct?</div>',
                    unsafe_allow_html=True,
                )
                _fb_col1, _fb_col2, _fb_col3, _ = st.columns([1, 1.4, 1, 4.6])
                with _fb_col1:
                    if st.button("Yes", key="fb_yes", use_container_width=True):
                        _fq = st.session_state.pop("_feedback_query", "")
                        st.session_state.learned_answers.append({
                            "query":   _fq,
                            "message": dict(msg),
                        })
                        _save_learned_to_disk(st.session_state.learned_answers)
                        st.session_state.pop("_awaiting_feedback", None)
                        st.toast("Saved! I'll use this answer for similar questions.", icon="✅")
                        st.rerun()
                with _fb_col2:
                    if st.button("Partially", key="fb_partial", use_container_width=True):
                        # Do NOT permanently save — just retry for a more complete answer.
                        # Only "Yes" persists to disk.
                        _fq = st.session_state.pop("_feedback_query", "")
                        st.session_state.pop("_awaiting_feedback", None)
                        if st.session_state.messages and st.session_state.messages[-1]["role"] == "assistant":
                            st.session_state.messages.pop()
                        st.session_state["_pending_query"] = _fq
                        st.session_state["_retry_mode"]    = "partial"  # keeps use_direct=True
                        st.toast("Looking for more complete information…", icon="🔍")
                        st.rerun()
                with _fb_col3:
                    if st.button("No", key="fb_no", use_container_width=True):
                        _fq = st.session_state.pop("_feedback_query", "")
                        st.session_state.pop("_awaiting_feedback", None)
                        if st.session_state.messages and st.session_state.messages[-1]["role"] == "assistant":
                            st.session_state.messages.pop()
                        st.session_state["_pending_query"] = _fq
                        st.session_state["_retry_mode"]    = True
                        st.rerun()

    # Cancel button — visible in chat area while processing
    if st.session_state.get("_pending_query"):
        _cc, _ = st.columns([2, 8])
        with _cc:
            st.markdown(
                '<div style="font-size:.8rem;color:#c8c5ff;padding:2px 0 4px;">⏳ Searching documents…</div>',
                unsafe_allow_html=True,
            )
            if st.button("⏹ Cancel Query", key="stop_btn", use_container_width=True,
                         help="Cancel this query"):
                st.session_state.pop("_pending_query", None)
                st.session_state.pop("_processing", None)
                st.toast("Query cancelled.", icon="🛑")
                st.rerun()

    # Thinking indicator — shown while query is pending
    thinking_slot = st.empty()
    if st.session_state.get("_pending_query"):
        thinking_slot.markdown("""
        <div class="msg-bot">
            <div class="avatar">🤖</div>
            <div class="bubble" style="padding:14px 20px;">
                <span class="thinking-dots">
                    <span></span><span></span><span></span>
                </span>
            </div>
        </div>
        """, unsafe_allow_html=True)

    # Unique anchor — ID changes with message count so browser always sees it as new
    _anchor_id = f"chat-bottom-{len(st.session_state.get('messages', []))}"
    st.markdown(f'<div id="{_anchor_id}" style="height:1px;"></div>', unsafe_allow_html=True)

    st.markdown("</div>", unsafe_allow_html=True)

    # Auto-scroll — scrollIntoView finds the real scrollable ancestor automatically
    st.markdown(f"""
    <script>
    (function() {{
        function scrollToAnchor() {{
            var el = document.getElementById('{_anchor_id}');
            if (el) {{
                el.scrollIntoView({{block: 'end'}});
            }}
        }}
        scrollToAnchor();
        setTimeout(scrollToAnchor, 150);
        setTimeout(scrollToAnchor, 500);
        setTimeout(scrollToAnchor, 1000);
    }})();
    </script>
    """, unsafe_allow_html=True)

    # ── Query form — only rendered when not processing ───────────────────────
    # Cancel Query button is shown in the chat area above, so no disabled
    # form is needed during processing. Rendering it causes a duplicate bar
    # during the rerun transition when the old disabled form and new enabled
    # form are both briefly visible.
    if not st.session_state.get("_pending_query"):
        with st.form(key="query_form", clear_on_submit=True):
            col_in, col_btn = st.columns([5, 1])
            with col_in:
                query = st.text_input(
                    "Ask a question",
                    placeholder="What does the document say about…?",
                    label_visibility="collapsed",
                )
            with col_btn:
                submit = st.form_submit_button("Send ➤", use_container_width=True)

        if submit and query.strip():
            st.session_state.messages.append({"role": "user", "content": query.strip()})
            st.session_state["_pending_query"] = query.strip()
            st.session_state.pop("_awaiting_feedback", None)
            st.rerun()

    # Process pending query (only runs when _pending_query is still set after the
    # Stop button check above — i.e., user did not click Stop on this rerun)
    if st.session_state.get("_pending_query"):
        from collections import Counter
        pending = st.session_state.pop("_pending_query")
        _retry_mode = st.session_state.pop("_retry_mode", False)

        # ── Check learned answers first ───────────────────────────────────────
        # If we previously confirmed a correct answer for a similar question,
        # replay it immediately without going to retrieval.
        # Guard: if both the pending query and the learned query name a CDISC
        # domain, they must name the SAME domain — otherwise "CM domain edit
        # checks" would incorrectly serve the cached "DM domain edit checks"
        # answer (both score ≥ 2 on shared words domain/edit/checks because
        # 2-char codes are invisible to _heading_match).
        _served_from_learned = False
        _pending_domain = _detect_domain(pending)
        for la in st.session_state.learned_answers:
            if _heading_match(pending, la["query"]) >= 2:
                _la_domain = _detect_domain(la.get("query", ""))
                if _pending_domain and _la_domain and _pending_domain != _la_domain:
                    continue   # different domain — do not serve this cached answer
                msg = dict(la["message"])   # shallow copy so we don't mutate stored
                st.session_state.messages.append(msg)
                st.session_state["_awaiting_feedback"] = True
                st.session_state["_feedback_query"]    = pending
                thinking_slot.empty()
                _served_from_learned = True
                st.rerun()
                break

        if _served_from_learned:
            st.stop()
        _is_table_q = bool(re.search(
            r'\b(assessment\s+schedule|schedule\s+of\s+assessments?|'
            r'visit\s+schedule|time\s+and\s+events?|'
            r'procedures?\s+and\s+assessments?)\b',
            pending, re.IGNORECASE
        ))
        _domain_in_q = _detect_domain(pending)
        relevant = retrieve(pending, st.session_state.doc_chunks, st.session_state.embeddings,
                            top_k=15 if _retry_mode else (12 if _domain_in_q else (8 if _is_table_q else 5)))

        # ── Hard document filter ──────────────────────────────────────────────
        # Re-run _doc_filter to find which documents the query refers to.
        # If specific documents were identified, remove any chunks from other docs
        # that may have crept in via table expansion or semantic search.
        _filter_pool, _ = _doc_filter(pending, st.session_state.doc_chunks)
        _matched_srcs = {c["source"] for c in _filter_pool}
        _all_srcs     = {c["source"] for c in st.session_state.doc_chunks}
        if _matched_srcs != _all_srcs:
            # A specific document was identified → enforce strictly
            relevant = [c for c in relevant if c["source"] in _matched_srcs]
            if not relevant:
                relevant = _filter_pool[:8]  # fallback if over-filtered

        # ── Domain filter (DM, VS, AE, etc.) ─────────────────────────────────
        # When user asks for a specific CDISC domain, keep only chunks that
        # mention that domain — prevents returning all domains from the file.
        _query_domain = _domain_in_q   # used for table-row filter below
        _domain_result = _domain_filter(pending, relevant)
        if _domain_result is not None:
            relevant = _domain_result or relevant

        context = build_context(relevant)

        # ── Image classification (table / relevant / irrelevant) ─────────────
        api_key = st.session_state.get("anthropic_api_key", "") or os.environ.get("ANTHROPIC_API_KEY", "")

        # Deduplicate: images appearing on multiple pages are logos → discard
        _img_cnt: Counter = Counter()
        for c in relevant:
            for b64 in c.get("images", []):
                _img_cnt[b64] += 1
        _uniq_imgs: list = []
        _seen_b64: set = set()
        for c in relevant:
            for b64 in c.get("images", []):
                if _img_cnt[b64] == 1 and b64 not in _seen_b64:
                    _uniq_imgs.append((b64, c["source"], c["page"]))
                    _seen_b64.add(b64)

        # ── Score each chunk by heading match ────────────────────────────────
        # Use max of extracted-heading score and text-scan score so eCRF-style
        # pages with repeated document headers are still found correctly.
        for c in relevant:
            c["_hscore"] = max(
                _heading_match(pending, c.get("heading", "")),
                _text_heading_match(pending, c.get("text", "")),
            )
        best_hscore = max((c["_hscore"] for c in relevant), default=0)

        # ── Route: direct render (structure-preserving) vs LLM ───────────────
        # Direct render when:
        #   • query is explicitly for a table/schedule (_is_table_q), OR
        #   • query matches a section heading with ≥2 meaningful words
        # In retry mode (user said "No"), force the LLM path regardless of heading match
        # "partial" retry keeps direct render (heading match still valid)
        # boolean True retry (full No) forces LLM for a fresh interpretation
        use_direct = (_retry_mode is not True) and (_is_table_q or best_hscore >= 2)

        if use_direct:
            # ── DIRECT RENDER — preserve document structure exactly ───────────
            # Select chunks: table queries use all relevant; heading queries use
            # only well-matched chunks, sorted best-match first.
            if _is_table_q:
                # For table queries: prefer chunks whose heading scores > 0
                # (near the schedule heading), fall back to all relevant only if needed.
                sched_chunks = [c for c in relevant if c["_hscore"] > 0]
                render_chunks = sched_chunks if sched_chunks else relevant
            else:
                render_chunks = sorted(
                    [c for c in relevant if c["_hscore"] >= 2],
                    key=lambda c: (-c["_hscore"], c["page"]),
                )
                # Include the page immediately after each matched heading so figures
                # that appear below section text (or on the next page) are captured.
                _rc_sp = {(c["source"], c["page"]) for c in render_chunks}
                _rel_by_sp2 = {(c["source"], c["page"]): c for c in relevant}
                for _rc in list(render_chunks):
                    _sp_next = (_rc["source"], _rc["page"] + 1)
                    if _sp_next not in _rc_sp and _sp_next in _rel_by_sp2:
                        render_chunks.append(_rel_by_sp2[_sp_next])
                        _rc_sp.add(_sp_next)

            # Classify each chunk by its dominant content type
            direct_tables:  list = []
            direct_images:  list = []   # (b64, src, pg) — page renders + classified images
            direct_texts:   list = []   # (text, src, pg) — verbatim text sections
            seen_pages: set = set()
            seen_img:   set = set()
            import html as _html_mod

            for c in render_chunks:
                src, pg = c["source"], c["page"]
                page_key = (src, pg)

                # ── Structured table (Method 1 / find_tables) → dataframe ──
                for t in c.get("tables", []):
                    if not _is_quality_table(t):
                        continue
                    # When a specific CDISC domain is requested, keep only rows
                    # for that domain — prevents showing the full multi-domain table.
                    if _query_domain:
                        t = _filter_table_rows_by_domain(t, _query_domain)
                        if len(t) < 2:   # header only → no matching rows, skip
                            continue
                    if page_key not in seen_pages:
                        seen_pages.add(page_key)
                    direct_tables.append((t, src, pg))

                # ── No structured table → page image only for table/figure queries ──
                # For plain-text sections (e.g. eCRF completion text), show the text
                # instead of rendering a page image that includes the document header.
                _FIGURE_RE = re.compile(r'\bFigure\s+\d', re.IGNORECASE)
                _has_fig_ref = bool(_FIGURE_RE.search(c.get("text", "")))
                _needs_img = _is_table_q or _has_fig_ref
                if _needs_img and page_key not in seen_pages and src.lower().endswith(".pdf"):
                    hint = c.get("heading", "") or pending
                    img_b64 = get_table_image(src, pg, heading_hint=hint)
                    if img_b64:
                        seen_pages.add(page_key)
                        direct_images.append((img_b64, src, pg))
                # Figure reference found → also render the NEXT page (figure usually
                # appears on the page after the text that cites it)
                if _has_fig_ref and src.lower().endswith(".pdf"):
                    _next_key = (src, pg + 1)
                    if _next_key not in seen_pages:
                        _nxt_img = get_table_image(src, pg + 1, heading_hint="Figure")
                        if _nxt_img:
                            seen_pages.add(_next_key)
                            direct_images.append((_nxt_img, src, pg + 1))

                # ── Embedded images in the document ──
                # Only shown for table/figure queries; for text-guideline queries
                # the embedded form screenshots are supplementary noise.
                if _is_table_q:
                    for b64 in c.get("images", []):
                        if b64 not in seen_img:
                            seen_img.add(b64)
                            if api_key and ANTHROPIC_OK:
                                kind, rows = analyze_image(b64, pending, api_key)
                                if kind == "table" and rows and _is_quality_table(rows):
                                    direct_tables.append((rows, src, pg))
                                elif kind == "relevant":
                                    direct_images.append((b64, src, pg))
                            else:
                                direct_images.append((b64, src, pg))

                # ── Text sections (PDF and non-PDF) ──
                # Show text when no quality table has claimed this page.
                # (seen_pages is only updated by quality tables and page images)
                if page_key not in seen_pages:
                    if c.get("text", "").strip():
                        seen_pages.add(page_key)
                        direct_texts.append((c["text"].strip(), src, pg))

            # Merge multi-page tables into one — applies to both domain queries and
            # schedule/table queries where the same table spans multiple pages.
            if (_query_domain or _is_table_q) and len(direct_tables) > 1:
                _merged_header = direct_tables[0][0][0]   # header row from first table
                _merged_rows = []
                _seen_row_keys: set = set()
                for _dt, _ds, _dp in direct_tables:
                    for _row in _dt[1:]:   # skip header of each table
                        _key = tuple(str(c).strip() for c in _row)
                        if _key not in _seen_row_keys:
                            _merged_rows.append(_row)
                            _seen_row_keys.add(_key)
                if _merged_rows:
                    direct_tables = [([_merged_header] + _merged_rows, direct_tables[0][1], direct_tables[0][2])]

            # Build the answer bubble — minimal heading + source info only
            # Always derive heading from the user's own query for domain queries:
            # chunk headings often reflect a different domain (e.g. "AE Domain
            # Edit Checks" stored in the doc when the user asked for CM).
            _q_domain_for_hdr = _query_domain or _detect_domain(pending)
            top_heading = ""
            if _q_domain_for_hdr:
                _th = pending.strip().title()
                top_heading = re.sub(
                    r'\b' + re.escape(_q_domain_for_hdr.title()) + r'\b',
                    _q_domain_for_hdr.upper(), _th
                )
            elif render_chunks:
                best = max(render_chunks, key=lambda c: c["_hscore"])
                top_heading = best.get("heading") or pending.strip().title()
            answer_html = f"<b>{_html_mod.escape(top_heading)}</b>"

            # Verbatim text (DOCX / TXT / non-PDF sources)
            if direct_texts:
                text_html = "<br><br>".join(
                    _html_mod.escape(txt).replace("\n", "<br>")
                    + f'<br><span style="font-size:.75rem;color:#9090b8;">'
                      f'[{_html_mod.escape(src)}, Page {pg}]</span>'
                    for txt, src, pg in direct_texts
                )
                answer_html += "<br><br>" + text_html

            thinking_slot.empty()
            sources = list({f"{c['source']} p.{c['page']}" for c in render_chunks})
            # Build copyable plain-text version: heading + any text sections + sources
            _copy_content = top_heading
            if direct_texts:
                _copy_content += "\n\n" + "\n\n".join(
                    f"{txt}\n[{src}, Page {pg}]" for txt, src, pg in direct_texts
                )
            else:
                _copy_content += "\n\n" + "\n".join(f"[{s}]" for s in sources)
            st.session_state.messages.append({
                "role": "assistant",
                "content": _copy_content,
                "answer_html": answer_html,
                "images": direct_images,
                "tables": direct_tables,
                "sources": sources,
            })
            st.session_state["_awaiting_feedback"] = True
            st.session_state["_feedback_query"]    = pending

        else:
            # ── LLM PATH — open-ended questions with no strong heading match ──
            display_images: list = []
            vision_images:  list = []
            img_tables:     list = []

            if api_key and ANTHROPIC_OK and _uniq_imgs:
                with st.spinner("Analysing images…"):
                    for b64, src, page in _uniq_imgs[:10]:
                        kind, rows = analyze_image(b64, pending, api_key)
                        if kind == "table" and rows:
                            img_tables.append((rows, src, page))
                            cols   = len(rows[0])
                            hdr    = "| " + " | ".join(str(x) for x in rows[0]) + " |"
                            sep_ln = "| " + " | ".join("---" for _ in range(cols)) + " |"
                            body   = "\n".join(
                                "| " + " | ".join(str(x) for x in r) + " |" for r in rows[1:]
                            )
                            context += (f"\n\n---\n\n[Image Table from {src}, Page {page}]"
                                        f"\n[TABLE]\n{hdr}\n{sep_ln}\n{body}\n[/TABLE]")
                        elif kind == "relevant":
                            display_images.append((b64, src, page))
                            vision_images.append((b64, src, page))
            else:
                display_images = _uniq_imgs

            # Render page images for any "Figure N" references found in retrieved text
            _FIGURE_RE_LLM = re.compile(r'\bFigure\s+\d', re.IGNORECASE)
            _fig_seen_pg: set = {(b, s, p) for b, s, p in display_images}
            _fig_seen_keys: set = set()
            for _fc in relevant:
                if _FIGURE_RE_LLM.search(_fc.get("text", "")) and _fc["source"].lower().endswith(".pdf"):
                    for _offset in (0, 1):   # render the citing page AND the next page
                        _fk = (_fc["source"], _fc["page"] + _offset)
                        if _fk not in _fig_seen_keys:
                            _fig_seen_keys.add(_fk)
                            _fi = get_table_image(_fc["source"], _fc["page"] + _offset, heading_hint="Figure")
                            if _fi and _fi not in {b for b, _, _ in display_images}:
                                display_images.append((_fi, _fc["source"], _fc["page"] + _offset))

            raw_answer = ask_llm(pending, context, vision_images=vision_images,
                                 table_query=_is_table_q)
            thinking_slot.empty()
            answer_html, _, extra_tables = render_answer(raw_answer, relevant)

            # Apply domain filter and dedup-merge for LLM-path tables
            if _query_domain and extra_tables:
                _dom_filtered = []
                for _et, _es, _ep in extra_tables:
                    _tf = _filter_table_rows_by_domain(_et, _query_domain)
                    if len(_tf) >= 2:
                        _dom_filtered.append((_tf, _es, _ep))
                if _dom_filtered:
                    extra_tables = _dom_filtered
                    if len(extra_tables) > 1:
                        _mh = extra_tables[0][0][0]
                        _mr, _sk = [], set()
                        for _et, _es, _ep in extra_tables:
                            for _row in _et[1:]:
                                _k = tuple(str(c).strip() for c in _row)
                                if _k not in _sk:
                                    _mr.append(_row)
                                    _sk.add(_k)
                        if _mr:
                            extra_tables = [([_mh] + _mr, extra_tables[0][1], extra_tables[0][2])]

            sources = list({f"{c['source']} p.{c['page']}" for c in relevant})
            st.session_state.messages.append({
                "role": "assistant",
                "content": raw_answer,
                "answer_html": answer_html,
                "images": display_images,
                "tables": extra_tables + img_tables,
                "sources": sources,
            })
            st.session_state["_awaiting_feedback"] = True
            st.session_state["_feedback_query"]    = pending

        st.rerun()

